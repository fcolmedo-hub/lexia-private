from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt


class DocxExporter:
    def export_markdown_like(
        self,
        title: str,
        content: str,
        destination: str | Path,
    ) -> Path:
        path = Path(destination)
        path.parent.mkdir(parents=True, exist_ok=True)

        document = DocxDocument()
        document.add_heading(title, level=0)

        for raw_line in content.splitlines():
            line = raw_line.rstrip()

            if not line:
                document.add_paragraph()
                continue

            if line.startswith("# "):
                document.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                document.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                document.add_heading(line[4:].strip(), level=3)
            elif line.startswith("- "):
                document.add_paragraph(line[2:].strip(), style="List Bullet")
            elif line[:3].rstrip(".").isdigit() and ". " in line:
                document.add_paragraph(
                    line.split(". ", 1)[1],
                    style="List Number",
                )
            else:
                paragraph = document.add_paragraph(line)
                paragraph.style.font.size = Pt(11)

        document.save(path)
        return path
